import os
import re
import json
import uuid
import logging
import tempfile
import shutil
import traceback
import platform
import subprocess
import signal
from datetime import datetime
from collections import defaultdict, deque
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io
import sys
try:
    from lxml import etree
except ImportError:
    import xml.etree.ElementTree as etree

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("docx_parser.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def iter_block_items(parent):
    """
    按文档顺序生成段落和表格，增强异常处理
    """
    try:
        if isinstance(parent, _Document):
            parent_elem = parent.element.body
        elif isinstance(parent, Table):
            parent_elem = parent._tbl
        else:
            logger.warning(f"不支持的父级类型: {type(parent)}")
            return
        
        if parent_elem is None:
            logger.warning("父级元素为空，跳过解析")
            return
            
        for child in parent_elem.iterchildren():
            try:
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)
            except Exception as e:
                logger.warning(f"解析文档块时遇到异常，跳过: {e}")
                continue
                
    except Exception as e:
        logger.error(f"遍历文档块失败: {e}")
        return

def clean_text(text):
    """清理文本中的特殊字符和多余空格"""
    if not text:
        return ""
    # 替换各种空格和特殊字符
    text = text.replace('\xa0', ' ').replace('\t', ' ').replace('\r', '')
    # 合并多余空格
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def get_heading_level(para):
    """
    智能识别标题层级
    1. 通过样式名称识别（如"Heading 1"或"标题1"）
    2. 通过文本模式识别（如"1.1 标题内容"）
    """
    # 通过样式名称识别
    if hasattr(para, 'style') and para.style and para.style.name:
        style_name = para.style.name.lower()
        if 'heading' in style_name or '标题' in style_name:
            # 提取样式中的数字
            match = re.search(r'\d+', style_name)
            if match:
                return int(match.group())
            return 1  # 默认一级标题
    
    # 通过文本模式识别 (如 "1.1 标题内容")
    text = clean_text(para.text)
    match = re.match(r'^(\d+(\.\d+)*)\s+', text)
    if match:
        level_str = match.group(1)
        return min(len(level_str.split('.')), 6)  # 最大支持6级
    
    return 0

def is_list_item(para):
    """判断段落是否为列表项"""
    if not hasattr(para, '_p') or para._p is None:
        return False
    if not hasattr(para._p, 'pPr') or para._p.pPr is None:
        return False
    return para._p.pPr.numPr is not None

def get_list_info(para, list_counter):
    """
    获取列表项详细信息
    返回: (list_level, prefix)
    """
    num_pr = para._p.pPr.numPr
    if num_pr is None:
        return 0, ""
    
    # 获取列表层级
    ilvl = num_pr.ilvl
    list_level = int(ilvl.val) if ilvl is not None and ilvl.val is not None else 0
    
    # 获取列表类型
    num_id = num_pr.numId
    num_id_val = num_id.val if num_id is not None else None
    
    # 更新列表计数器
    list_counter[list_level] += 1
    # 重置子层级计数器
    for lvl in range(list_level + 1, 10):
        if lvl in list_counter:
            list_counter[lvl] = 0
    
    # 创建前缀
    prefix = " " * (list_level * 4)  # 每级缩进4个空格
    
    # 确定列表符号
    is_bullet = "bullet" in str(para.style.name).lower()
    if is_bullet:
        prefix += "• "  # 项目符号
    else:
        # 数字编号
        prefix += ".".join(str(list_counter[lvl]) for lvl in range(list_level + 1)) + ". "
    
    return list_level, prefix

def extract_smartart_from_xml(xml_str, doc_part, output_dir, context=""):
    """
    从XML字符串中提取SmartArt图表信息
    返回: SmartArt节点列表
    """
    smartart_nodes = []
    try:
        if not xml_str or ('<a:graphic' not in xml_str and '<w:drawing>' not in xml_str):
            return smartart_nodes
        
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'dgm': 'http://schemas.openxmlformats.org/drawingml/2006/diagram',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006'
        }
        
        root = etree.fromstring(xml_str)
        
        # 查找所有graphic元素
        graphics = root.findall('.//a:graphic', namespaces)
        
        for graphic in graphics:
            graphic_data = graphic.find('.//a:graphicData', namespaces)
            if graphic_data is not None:
                uri = graphic_data.get('uri', '')
                
                # 检查是否为SmartArt (diagram)
                if 'diagram' in uri:
                    logger.info(f"发现SmartArt图表在 {context}")
                    
                    # 提取SmartArt详细信息
                    smartart_info = extract_smartart_details(graphic_data, doc_part, output_dir, namespaces)
                    if smartart_info:
                        smartart_info["context"] = context
                        smartart_nodes.append(smartart_info)
    
    except Exception as e:
        logger.error(f"从XML提取SmartArt失败: {e}")
    
    return smartart_nodes

def extract_smartart_details(graphic_data, doc_part, output_dir, namespaces):
    """
    提取SmartArt的详细信息和文本内容
    """
    try:
        # 查找关系ID
        rel_ids_elem = graphic_data.find('.//dgm:relIds', namespaces)
        if rel_ids_elem is None:
            logger.warning("未找到SmartArt关系ID元素")
            return None
        
        # 获取各种关系ID（使用完整的命名空间属性名）
        dm_rel_id = rel_ids_elem.get(f'{{{namespaces["r"]}}}dm')  # data model
        lo_rel_id = rel_ids_elem.get(f'{{{namespaces["r"]}}}lo')  # layout
        qs_rel_id = rel_ids_elem.get(f'{{{namespaces["r"]}}}qs')  # quick style
        cs_rel_id = rel_ids_elem.get(f'{{{namespaces["r"]}}}cs')  # color scheme
        
        logger.info(f"SmartArt关系ID - dm: {dm_rel_id}, lo: {lo_rel_id}, qs: {qs_rel_id}, cs: {cs_rel_id}")
        logger.info(f"可用的关系部件数量: {len(doc_part.related_parts)}")
        
        smartart_data = {
            "type": "smartart",
            "text_content": [],
            "diagram_type": "unknown",
            "nodes": []
        }
        
        # 提取数据模型中的文本内容
        if dm_rel_id and dm_rel_id in doc_part.related_parts:
            logger.info(f"找到数据模型关系: {dm_rel_id}")
            data_part = doc_part.related_parts[dm_rel_id]
            data_xml = data_part.blob.decode('utf-8')
            text_content = extract_smartart_text(data_xml)
            smartart_data["text_content"] = text_content
            smartart_data["nodes"] = text_content  # 兼容性字段
            logger.info(f"提取到 {len(text_content)} 个文本节点")
        else:
            logger.warning(f"未找到数据模型关系 {dm_rel_id} 或关系不存在")
        
        # 尝试确定图表类型
        if lo_rel_id and lo_rel_id in doc_part.related_parts:
            layout_part = doc_part.related_parts[lo_rel_id]
            layout_xml = layout_part.blob.decode('utf-8')
            diagram_type = extract_diagram_type(layout_xml)
            if diagram_type:
                smartart_data["diagram_type"] = diagram_type
        
        # 生成唯一ID
        smartart_id = f"smartart_{uuid.uuid4().hex[:8]}"
        smartart_data["id"] = smartart_id
        
        # 保存原始数据到文件（可选）
        smartart_dir = os.path.join(output_dir, "smartart")
        os.makedirs(smartart_dir, exist_ok=True)
        
        smartart_file = os.path.join(smartart_dir, f"{smartart_id}.json")
        with open(smartart_file, 'w', encoding='utf-8') as f:
            json.dump(smartart_data, f, ensure_ascii=False, indent=2)
        
        smartart_data["file_path"] = f"smartart/{smartart_id}.json"
        
        logger.info(f"提取SmartArt成功，包含 {len(smartart_data['text_content'])} 个文本节点")
        return smartart_data
        
    except Exception as e:
        logger.error(f"提取SmartArt详细信息失败: {e}")
        logger.error(traceback.format_exc())
        return None

def extract_smartart_text(data_xml):
    """
    从SmartArt数据模型XML中提取文本内容，保持层次结构
    """
    text_nodes = []
    try:
        # 定义命名空间
        namespaces = {
            'dgm': 'http://schemas.openxmlformats.org/drawingml/2006/diagram',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
        }
        
        root = etree.fromstring(data_xml.encode('utf-8'))
        
        # 1. 解析所有数据点
        points = {}
        pt_elements = root.findall('.//dgm:pt', namespaces)
        
        for pt in pt_elements:
            model_id = pt.get('modelId')
            point_type = pt.get('type', 'unknown')
            
            # 提取文本内容
            text_parts = []
            text_elems = pt.findall('.//a:t', namespaces)
            for text_elem in text_elems:
                if text_elem.text and text_elem.text.strip():
                    text_parts.append(text_elem.text.strip())
            
            combined_text = ' '.join(text_parts) if text_parts else ''
            points[model_id] = {
                'text': combined_text,
                'type': point_type,
                'children': []
            }
        
        # 2. 解析连接关系构建层次结构
        connections = root.findall('.//dgm:cxn', namespaces)
        parent_child_map = {}
        
        for cxn in connections:
            src_id = cxn.get('srcId')
            dest_id = cxn.get('destId')
            cxn_type = cxn.get('type', 'unknown')
            src_ord = int(cxn.get('srcOrd', 0))
            
            # 只处理非presentation的连接（即真正的数据结构关系）
            if cxn_type in ['presOf', 'presParOf']:
                continue
                
            if src_id not in parent_child_map:
                parent_child_map[src_id] = []
            parent_child_map[src_id].append((dest_id, src_ord))
        
        # 3. 找到根节点（type="doc"的节点）
        root_id = None
        for point_id, point in points.items():
            if point['type'] == 'doc':
                root_id = point_id
                break
        
        # 4. 构建层次结构并提取文本
        def extract_hierarchy_text(node_id, level=0):
            if node_id not in points:
                return []
            
            texts = []
            point = points[node_id]
            
            # 添加当前节点的文本（如果有的话）
            if point['text']:
                text_nodes.append({
                    "text": point['text'],
                    "type": "text_node", 
                    "level": level,
                    "node_id": node_id
                })
                texts.append(point['text'])
            
            # 处理子节点，按srcOrd排序
            if node_id in parent_child_map:
                children = sorted(parent_child_map[node_id], key=lambda x: x[1])
                for child_id, _ in children:
                    child_texts = extract_hierarchy_text(child_id, level + 1)
                    texts.extend(child_texts)
            
            return texts
        
        if root_id:
            extract_hierarchy_text(root_id)
        
        # 如果层次结构提取失败，回退到简单文本提取
        if not text_nodes:
            seen_texts = set()
            for point in points.values():
                if point['text'] and point['text'] not in seen_texts:
                    seen_texts.add(point['text'])
                    text_nodes.append({
                        "text": point['text'],
                        "type": "text_node"
                    })
    
    except Exception as e:
        logger.error(f"提取SmartArt文本失败: {e}")
        # 回退到原始方法
        try:
            root = etree.fromstring(data_xml.encode('utf-8'))
            texts = []
            # 使用兼容的方法查找文本元素
            text_elements = []
            try:
                # 尝试使用 lxml 的 xpath
                xpath_method = getattr(root, 'xpath', None)
                if xpath_method:
                    text_elements = xpath_method('.//*[local-name()="t"]')
                else:
                    raise AttributeError("xpath not available")
            except (AttributeError, ImportError):
                # 回退到遍历方法
                for elem in root.iter():
                    if elem.tag.endswith('}t') or elem.tag == 't':
                        text_elements.append(elem)
            
            seen_texts = set()
            for elem in text_elements:
                if elem.text and elem.text.strip():
                    text_content = elem.text.strip()
                    if text_content not in seen_texts:
                        seen_texts.add(text_content)
                        text_nodes.append({
                            "text": text_content,
                            "type": "text_node"
                        })
        except:
            pass
    
    return text_nodes

def extract_diagram_type(layout_xml):
    """
    从布局XML中确定图表类型
    """
    try:
        # 常见的SmartArt布局类型
        layout_patterns = {
            'list': ['list', 'bullet', 'sequence'],
            'process': ['process', 'flow', 'step'],
            'cycle': ['cycle', 'circular'],
            'hierarchy': ['hierarchy', 'org', 'tree'],
            'relationship': ['relationship', 'venn', 'matrix'],
            'pyramid': ['pyramid', 'funnel']
        }
        
        layout_xml_lower = layout_xml.lower()
        
        for diagram_type, patterns in layout_patterns.items():
            for pattern in patterns:
                if pattern in layout_xml_lower:
                    return diagram_type
        
        return "unknown"
    
    except Exception as e:
        logger.error(f"确定图表类型失败: {e}")
        return "unknown"

def extract_embedded_objects_from_xml(xml_str, doc_part, output_dir, context="", quick_mode=True):
    """
    从XML字符串中提取嵌入对象（如Visio图表、Excel表格等）
    返回: 嵌入对象节点列表
    """
    embedded_objects = []
    try:
        if not xml_str or '<w:object' not in xml_str:
            return embedded_objects
        
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'o': 'urn:schemas-microsoft-com:office:office',
            'v': 'urn:schemas-microsoft-com:vml',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        root = etree.fromstring(xml_str)
        
        # 查找所有object元素
        objects = root.findall('.//w:object', namespaces)
        
        for obj_idx, obj in enumerate(objects):
            logger.info(f"发现嵌入对象在 {context}")
            
            # 分析OLE对象
            ole_objects = obj.findall('.//o:OLEObject', namespaces)
            
            for ole in ole_objects:
                ole_type = ole.get('Type', 'Unknown')
                prog_id = ole.get('ProgID', 'Unknown')
                r_id = ole.get(f'{{{namespaces["r"]}}}id')
                
                # 获取尺寸信息
                shapes = obj.findall('.//v:shape', namespaces)
                width, height = "Unknown", "Unknown"
                preview_image_r_id = None
                
                if shapes:
                    style = shapes[0].get('style', '')
                    width_match = re.search(r'width:([^;]+)', style)
                    height_match = re.search(r'height:([^;]+)', style)
                    if width_match:
                        width = width_match.group(1)
                    if height_match:
                        height = height_match.group(1)
                    
                    # 查找预览图像的关系ID
                    imagedatas = shapes[0].findall('.//v:imagedata', namespaces)
                    if imagedatas:
                        preview_image_r_id = imagedatas[0].get(f'{{{namespaces["r"]}}}id')
                
                # 确定对象类型
                object_type = "unknown"
                object_description = prog_id
                
                if "Visio" in prog_id:
                    object_type = "visio_diagram"
                    object_description = "Visio绘图/流程图"
                elif "Excel" in prog_id:
                    object_type = "excel_spreadsheet" 
                    object_description = "Excel电子表格"
                elif "PowerPoint" in prog_id:
                    object_type = "powerpoint_presentation"
                    object_description = "PowerPoint演示文稿"
                elif "Word" in prog_id:
                    object_type = "word_document"
                    object_description = "Word文档"
                
                # 尝试获取嵌入文件信息
                embedded_file_info = None
                if r_id and r_id in doc_part.related_parts:
                    try:
                        related_part = doc_part.related_parts[r_id]
                        embedded_file_info = {
                            "size": f"{len(related_part.blob)/1024:.2f} KB",
                            "content_type": getattr(related_part, 'content_type', 'Unknown')
                        }
                    except Exception as e:
                        logger.warning(f"无法获取嵌入文件信息: {e}")
                
                # 生成唯一ID
                object_id = f"embedded_obj_{uuid.uuid4().hex[:8]}"
                
                # 创建嵌入对象节点
                embedded_obj = {
                    "type": "embedded_object",
                    "object_type": object_type,
                    "description": object_description,
                    "prog_id": prog_id,
                    "ole_type": ole_type,
                    "width": width,
                    "height": height,
                    "context": context,
                    "id": object_id
                }
                
                if embedded_file_info:
                    embedded_obj.update(embedded_file_info)
                
                # 提取并保存预览图像
                preview_image_path = None
                if preview_image_r_id and preview_image_r_id in doc_part.related_parts:
                    try:
                        preview_image_path = extract_preview_image(
                            doc_part.related_parts[preview_image_r_id],
                            output_dir,
                            object_id,
                            quick_mode
                        )
                        if preview_image_path:
                            embedded_obj["preview_image"] = preview_image_path
                            logger.info(f"成功提取预览图像: {preview_image_path}")
                    except Exception as e:
                        logger.warning(f"提取预览图像失败: {e}")
                
                # 保存对象信息到文件
                objects_dir = os.path.join(output_dir, "embedded_objects")
                os.makedirs(objects_dir, exist_ok=True)
                
                object_file = os.path.join(objects_dir, f"{object_id}.json")
                with open(object_file, 'w', encoding='utf-8') as f:
                    json.dump(embedded_obj, f, ensure_ascii=False, indent=2)
                
                embedded_obj["file_path"] = f"embedded_objects/{object_id}.json"
                
                embedded_objects.append(embedded_obj)
                logger.info(f"提取嵌入对象成功: {object_description} ({width} x {height})")
    
    except Exception as e:
        logger.error(f"从XML提取嵌入对象失败: {e}")
    
    return embedded_objects

def extract_preview_image(image_part, output_dir, object_id, quick_mode=True):
    """
    提取嵌入对象的预览图像并保存为文件
    """
    try:
        # 获取图像数据
        image_data = image_part.blob
        
        # 确定图像格式
        img_format = "png"  # 默认格式
        if hasattr(image_part, 'content_type'):
            content_type = image_part.content_type.lower()
            if 'jpeg' in content_type or 'jpg' in content_type:
                img_format = "jpg"
            elif 'gif' in content_type:
                img_format = "gif"
            elif 'bmp' in content_type:
                img_format = "bmp"
            elif 'png' in content_type:
                img_format = "png"
            elif 'emf' in content_type:
                img_format = "emf"
            elif 'wmf' in content_type:
                img_format = "wmf"
            elif 'tiff' in content_type:
                img_format = "tiff"
        
        # 如果是EMF/WMF格式，尝试转换为PNG
        if img_format in ['emf', 'wmf']:
            try:
                converted_path = convert_emf_to_png(image_data, output_dir, object_id, quick_mode=quick_mode)
                if converted_path:
                    return converted_path
            except Exception as e:
                logger.warning(f"EMF/WMF转换失败，保存原格式: {e}")
        
        # 创建images目录
        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        # 生成文件名
        img_filename = f"embedded_preview_{object_id}.{img_format}"
        image_path = os.path.join(images_dir, img_filename)
        
        # 保存图像
        with open(image_path, "wb") as img_file:
            img_file.write(image_data)
        
        # 返回相对路径
        return f"images/{img_filename}"
        
    except Exception as e:
        logger.error(f"提取预览图像失败: {e}")
        return None

def convert_emf_to_png(emf_data, output_dir, object_id, quick_mode=True):
    """
    自动将EMF格式转换为PNG格式，增强错误处理和健壮性
    quick_mode: 快速模式，跳过耗时的转换尝试，直接保存原格式
    """
    temp_emf_path = None
    try:
        # 验证输入数据
        if not emf_data or len(emf_data) < 16:
            logger.warning(f"EMF数据无效或太小: {len(emf_data) if emf_data else 0} bytes")
            return None
        
        # 创建images目录
        try:
            images_dir = os.path.join(output_dir, "images")
            os.makedirs(images_dir, exist_ok=True)
        except (OSError, IOError) as e:
            logger.error(f"创建images目录失败: {e}")
            return None
        
        # 快速模式：直接保存原格式，跳过转换
        if quick_mode:
            logger.info("快速模式：直接保存EMF文件，跳过转换")
            emf_filename = f"embedded_preview_{object_id}.emf"
            final_emf_path = os.path.join(images_dir, emf_filename)
            
            try:
                with open(final_emf_path, "wb") as f:
                    f.write(emf_data)
                logger.info(f"EMF文件已保存: {emf_filename}")
                return f"images/{emf_filename}"
            except Exception as e:
                logger.error(f"保存EMF文件失败: {e}")
                return None
        
        # 以下是原有的转换逻辑（在非快速模式下执行）
        # 创建临时EMF文件
        try:
            temp_emf_path = os.path.join(output_dir, f"temp_{object_id}.emf")
            with open(temp_emf_path, "wb") as f:
                f.write(emf_data)
            
            # 验证文件写入成功
            if not os.path.exists(temp_emf_path) or os.path.getsize(temp_emf_path) == 0:
                logger.error("临时EMF文件创建失败")
                return None
                
        except (OSError, IOError) as e:
            logger.error(f"写入临时EMF文件失败: {e}")
            return None
        
        # 输出PNG路径
        png_filename = f"embedded_preview_{object_id}.png"
        png_path = os.path.join(images_dir, png_filename)
        
        # 方法1: 使用 PIL 读取基本信息并尝试转换（限时5秒）
        use_timeout = False
        try:
            from PIL import Image
            
            # 只在Unix系统上使用信号超时
            use_timeout = hasattr(signal, 'SIGALRM')
            
            def timeout_handler(signum, frame):
                raise TimeoutError("PIL转换超时")
            
            # 设置5秒超时（仅在支持的系统上）
            if use_timeout:
                signal.signal(signal.SIGALRM, timeout_handler)
                signal.alarm(5)
            
            try:
                with Image.open(temp_emf_path) as img:
                    logger.info(f"PIL检测到图像: 模式={img.mode}, 尺寸={img.size}, 格式={img.format}")
                    
                    # 验证图像
                    img.verify()
                    # 重新打开进行转换
                    with Image.open(temp_emf_path) as img_convert:
                        # 尝试转换为RGB模式
                        if img_convert.mode in ('RGBA', 'LA'):
                            # 创建白色背景
                            background = Image.new('RGB', img_convert.size, (255, 255, 255))
                            if img_convert.mode == 'RGBA':
                                background.paste(img_convert, mask=img_convert.split()[-1])
                            else:
                                background.paste(img_convert)
                            img_convert = background
                        elif img_convert.mode != 'RGB':
                            img_convert = img_convert.convert('RGB')
                        
                        # 尝试保存为PNG
                        img_convert.save(png_path, 'PNG', optimize=True, quality=95)
                        
                        if os.path.exists(png_path) and os.path.getsize(png_path) > 0:
                            logger.info(f"使用 PIL 成功转换EMF为PNG: {png_filename}")
                            if temp_emf_path and os.path.exists(temp_emf_path):
                                os.remove(temp_emf_path)
                            return f"images/{png_filename}"
            finally:
                if use_timeout:
                    signal.alarm(0)  # 取消超时
                            
        except (TimeoutError, Exception) as e:
            logger.debug(f"PIL 转换失败或超时: {e}")
            if use_timeout:
                signal.alarm(0)  # 确保取消超时
        
        # 方法2: 使用系统工具转换（仅限macOS，限时10秒）
        if platform.system() == "Darwin":
            try:
                # 尝试使用sips命令（限时5秒）
                try:
                    logger.debug(f"尝试sips转换: {temp_emf_path}")
                    result = subprocess.run([
                        'sips', '-s', 'format', 'png', temp_emf_path, '--out', png_path
                    ], capture_output=True, text=True, timeout=5)
                    
                    if result.returncode == 0 and os.path.exists(png_path) and os.path.getsize(png_path) > 0:
                        logger.info(f"使用sips成功转换EMF为PNG: {png_filename}")
                        if temp_emf_path and os.path.exists(temp_emf_path):
                            os.remove(temp_emf_path)
                        return f"images/{png_filename}"
                    else:
                        logger.debug(f"sips转换失败: return code {result.returncode}")
                        
                except subprocess.TimeoutExpired:
                    logger.debug("sips命令超时")
                except FileNotFoundError:
                    logger.debug("sips命令不可用")
                except Exception as e:
                    logger.debug(f"sips转换异常: {e}")
            
            except Exception as e:
                logger.debug(f"系统命令转换失败: {e}")
        
        # 方法3: 转换失败，保存原始EMF文件
        try:
            logger.info("转换失败，保存原始EMF文件")
            
            # 保存原始EMF文件
            emf_filename = f"embedded_preview_{object_id}.emf"
            final_emf_path = os.path.join(images_dir, emf_filename)
            
            if temp_emf_path and os.path.exists(temp_emf_path):
                try:
                    os.rename(temp_emf_path, final_emf_path)
                    temp_emf_path = None  # 防止重复删除
                except OSError as e:
                    logger.warning(f"移动EMF文件失败，尝试复制: {e}")
                    try:
                        shutil.copy2(temp_emf_path, final_emf_path)
                    except Exception as copy_e:
                        logger.error(f"复制EMF文件也失败: {copy_e}")
                        return None
            
            logger.info(f"EMF文件已保存: {emf_filename}")
            
            # 返回EMF文件路径
            return f"images/{emf_filename}"
            
        except Exception as e:
            logger.error(f"保存EMF文件失败: {e}")
            return None
            
    except Exception as e:
        logger.error(f"EMF转换过程失败: {e}")
        return None
        
    finally:
        # 清理临时文件
        if temp_emf_path and os.path.exists(temp_emf_path):
            try:
                os.remove(temp_emf_path)
            except Exception as e:
                logger.debug(f"清理临时文件失败: {e}")

def extract_images_from_xml(xml_str, doc_part, images_dir, image_references, context="", quick_mode=True):
    """
    从XML字符串中提取图片并保存
    返回: 图片节点列表
    """
    image_nodes = []
    try:
        if not xml_str or ('<pic:pic' not in xml_str and '<w:drawing>' not in xml_str):
            return image_nodes
        
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }
        
        root = etree.fromstring(xml_str)
        
        # 查找所有图片元素
        for blip in root.findall('.//a:blip', namespaces):
            embed_id = blip.get(f'{{{namespaces["r"]}}}embed')
            if not embed_id:
                continue
                
            # 获取图片部件
            if embed_id not in doc_part.related_parts:
                logger.warning(f"图片关系 {embed_id} 在 {context} 中未找到")
                continue
                
            image_part = doc_part.related_parts[embed_id]
            image_data = image_part.blob
            
            # 确定图片格式
            img_format = "png"  # 默认格式
            if hasattr(image_part, 'content_type'):
                if 'jpeg' in image_part.content_type:
                    img_format = "jpg"
                elif 'gif' in image_part.content_type:
                    img_format = "gif"
                elif 'bmp' in image_part.content_type:
                    img_format = "bmp"
                elif 'png' in image_part.content_type:
                    img_format = "png"
                elif 'svg' in image_part.content_type:
                    img_format = "svg"
                elif 'tiff' in image_part.content_type:
                    img_format = "tiff"
            
            # 生成唯一图片ID
            image_id = f"img_{uuid.uuid4().hex[:8]}"
            img_filename = f"{image_id}.{img_format}"
            image_path = os.path.join(images_dir, img_filename)
            
            # 保存图片
            with open(image_path, "wb") as img_file:
                img_file.write(image_data)
            
            # 获取图片尺寸
            width, height = 0, 0
            try:
                with Image.open(io.BytesIO(image_data)) as img:
                    width, height = img.size
            except Exception as e:
                # 对于SVG等无法直接获取尺寸的图片格式，跳过尺寸获取
                pass
            
            # 创建图片节点
            image_node = {
                "type": "image",
                "url": f"images/{img_filename}",
                "format": img_format,
                "width": width,
                "height": height,
                "size": f"{len(image_data)/1024:.2f} KB",
                "context": context
            }
            
            # 记录图片信息
            image_references[image_id] = image_node
            image_nodes.append(image_node)
    
    except Exception as e:
        logger.error(f"从XML提取图片失败: {e}")
    
    return image_nodes

def extract_paragraph_content(para, output_dir, image_references, quick_mode=True):
    """提取段落中的图片、SmartArt和嵌入对象内容并返回节点列表"""
    content_nodes = []
    context = f"段落: {clean_text(para.text[:20])}..." if para.text else "段落"
    images_dir = os.path.join(output_dir, "images")
    
    for run_idx, run in enumerate(para.runs):
        if run._element is not None and run._element.xml:
            try:
                # 提取该运行中的图片
                images = extract_images_from_xml(
                    run._element.xml, 
                    para.part, 
                    images_dir, 
                    image_references,
                    f"{context} (运行 {run_idx})",
                    quick_mode
                )
                if images:
                    content_nodes.extend(images)
                
                # 提取SmartArt
                smartarts = extract_smartart_from_xml(
                    run._element.xml,
                    para.part,
                    output_dir,
                    f"{context} (运行 {run_idx})"
                )
                if smartarts:
                    content_nodes.extend(smartarts)
                
                # 提取嵌入对象 (OLE Objects)
                embedded_objects = extract_embedded_objects_from_xml(
                    run._element.xml,
                    para.part,
                    output_dir,
                    f"{context} (运行 {run_idx})",
                    quick_mode
                )
                if embedded_objects:
                    content_nodes.extend(embedded_objects)
                    
            except Exception as e:
                logger.error(f"提取段落内容失败: {e}")
    
    return content_nodes

def extract_paragraph_images(para, images_dir, image_references):
    """提取段落中的图片并返回图片节点列表（保持向后兼容）
    注意：此函数已被 extract_paragraph_content 替代，建议使用新函数
    """
    # 为了向后兼容，将 images_dir 转换为 output_dir 格式
    output_dir = os.path.dirname(images_dir) if images_dir.endswith('images') else images_dir
    content_nodes = extract_paragraph_content(para, output_dir, image_references, quick_mode=True)
    
    # 只返回图片节点，过滤其他类型
    return [node for node in content_nodes if node.get("type") == "image"]

def extract_table_images(cell, table_idx, row_idx, cell_idx, image_references, images_dir):
    """提取表格单元格中的图片"""
    image_nodes = []
    context = f"表格{table_idx}单元格[{row_idx},{cell_idx}]"
    
    try:
        if not hasattr(cell, '_tc') or cell._tc is None:
            return image_nodes
            
        cell_xml = etree.tostring(cell._tc, encoding='unicode')
        image_nodes = extract_images_from_xml(
            cell_xml, 
            cell.part, 
            images_dir, 
            image_references,
            context,
            quick_mode=True
        )
    except Exception as e:
        logger.error(f"提取表格图片失败: {e}")
    
    return image_nodes

def identify_merged_cells(table):
    """
    识别合并的单元格
    返回: 被合并单元格的位置集合{(row_idx, col_idx)}
    """
    merged_cells = set()
    try:
        grid = table._tbl.find('.//w:tblGrid', table._tbl.nsmap)
        if grid is None:
            return merged_cells
        
        cols = grid.findall('.//w:gridCol', grid.nsmap)
        if not cols:
            return merged_cells
        
        # 创建单元格网格
        rows = table.rows
        if not rows:
            return merged_cells
        
        # 标记所有被合并的单元格
        for row_idx, row in enumerate(rows):
            for cell_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.tcPr
                
                # 检测跨列
                grid_span = 1
                if tcPr.gridSpan is not None:
                    grid_span = int(tcPr.gridSpan.val)
                
                # 检测跨行
                row_span = 1
                v_merge = tcPr.vMerge
                if v_merge is not None and v_merge.val == "restart":
                    # 计算实际行跨度
                    for i in range(row_idx + 1, len(table.rows)):
                        if cell_idx >= len(table.rows[i].cells):
                            break
                        next_cell = table.rows[i].cells[cell_idx]
                        next_tcPr = next_cell._tc.tcPr
                        if next_tcPr.vMerge is None or next_tcPr.vMerge.val != "continue":
                            break
                        row_span += 1
                
                # 标记被合并的单元格
                if grid_span > 1 or row_span > 1:
                    for r in range(row_idx, row_idx + row_span):
                        for c in range(cell_idx, cell_idx + grid_span):
                            if r != row_idx or c != cell_idx:  # 跳过起始单元格
                                if r < len(rows) and c < len(rows[r].cells):
                                    merged_cells.add((r, c))
    except Exception as e:
        logger.error(f"识别合并单元格失败: {e}")
    
    return merged_cells

def parse_table(table, table_idx, image_references, images_dir):
    """解析表格并处理合并单元格"""
    # 识别所有合并单元格
    merged_cells = identify_merged_cells(table)
    
    table_data = []
    for row_idx, row in enumerate(table.rows):
        row_nodes = []
        for cell_idx, cell in enumerate(row.cells):
            # 如果是被合并的单元格，跳过
            if (row_idx, cell_idx) in merged_cells:
                continue
                
            # 创建单元格节点
            cell_node = {
                "type": "table_cell",
                "row": row_idx,
                "col": cell_idx,
                "content": []
            }
            
            # 添加文本内容
            cell_text = clean_text(cell.text)
            if cell_text:
                cell_node["content"].append({
                    "type": "text",
                    "text": cell_text
                })
                
            # 提取单元格中的图片
            image_nodes = extract_table_images(cell, table_idx, row_idx, cell_idx, image_references, images_dir)
            if image_nodes:
                for img in image_nodes:
                    cell_node["content"].append({
                        "type": "image",
                        "url": img["url"],
                        "format": img["format"],
                        "width": img["width"],
                        "height": img["height"],
                        "size": img["size"]
                    })
            
            row_nodes.append(cell_node)
        table_data.append(row_nodes)
    
    return table_data

def extract_metadata(doc, docx_path):
    """提取文档元数据"""
    core_props = doc.core_properties
    return {
        "source_path": docx_path,
        "title": core_props.title,
        "author": core_props.author,
        "created": str(core_props.created),
        "modified": str(core_props.modified),
        "subject": core_props.subject,
        "keywords": core_props.keywords,
        "category": core_props.category,
        "comments": core_props.comments,
        "company": getattr(core_props, "company", "N/A"),
        "file_size": f"{os.path.getsize(docx_path)/1024:.2f} KB"
    }

def extract_header_footer_images(doc, images_dir, image_references, quick_mode=True):
    """提取页眉页脚中的图片"""
    try:
        for section in doc.sections:
            # 页眉
            if hasattr(section, 'header') and section.header:
                for paragraph in section.header.paragraphs:
                    images = extract_paragraph_images(paragraph, images_dir, image_references)
                    if images:
                        # 简单记录页眉图片，不依赖 header_part 属性
                        logger.info(f"在页眉中找到 {len(images)} 张图片")
            
            # 页脚
            if hasattr(section, 'footer') and section.footer:
                for paragraph in section.footer.paragraphs:
                    images = extract_paragraph_images(paragraph, images_dir, image_references)
                    if images:
                        # 简单记录页脚图片，不依赖 footer_part 属性
                        logger.info(f"在页脚中找到 {len(images)} 张图片")
    except Exception as e:
        logger.error(f"提取页眉页脚图片失败: {e}")

def parse_docx(docx_path, output_dir, quick_mode=True):
    """
    解析单个DOCX文档并提取内容，增强错误处理和健壮性
    返回结构化JSON数据
    
    Args:
        docx_path: DOCX文件路径
        output_dir: 输出目录
        quick_mode: 快速模式，跳过耗时的EMF/WMF转换（默认True）
    """
    temp_dir = None
    try:
        # 首先检查输入文件
        if not os.path.exists(docx_path):
            logger.error(f"输入文件不存在: {docx_path}")
            return None
            
        if not docx_path.lower().endswith('.docx'):
            logger.error(f"不是有效的DOCX文件: {docx_path}")
            return None
            
        # 检查文件是否为临时文件或系统文件（以~$开头）
        filename = os.path.basename(docx_path)
        if filename.startswith('~$'):
            logger.warning(f"跳过临时文件: {filename}")
            return None
            
        file_size = os.path.getsize(docx_path)
        if file_size == 0:
            logger.error(f"文件为空: {docx_path}")
            return None
        elif file_size < 1024:  # 小于1KB的DOCX文件可能损坏
            logger.warning(f"文件可能损坏（太小）: {docx_path} ({file_size} bytes)")
        
        # 创建临时工作目录
        temp_dir = tempfile.mkdtemp(prefix="docx_extract_")
        logger.info(f"创建临时目录: {temp_dir}")
        
        # 复制文件到临时目录
        temp_docx_path = os.path.join(temp_dir, os.path.basename(docx_path))
        try:
            shutil.copy2(docx_path, temp_docx_path)
        except (OSError, IOError) as e:
            logger.error(f"复制文件失败: {e}")
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            return None
        
        # 检查文件是否成功复制
        if not os.path.exists(temp_docx_path) or os.path.getsize(temp_docx_path) == 0:
            logger.error(f"复制后的文件不存在或为空: {temp_docx_path}")
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            return None
        
        # 尝试打开文档
        try:
            doc = Document(temp_docx_path)
        except Exception as e:
            logger.error(f"无法打开DOCX文件 {docx_path}: {e}")
            if "Package not found" in str(e) or "not a valid" in str(e).lower():
                logger.error("文件可能损坏或不是有效的DOCX格式")
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            return None
        
        # 创建输出目录
        try:
            os.makedirs(output_dir, exist_ok=True)
            # 创建图片目录
            images_dir = os.path.join(output_dir, "images")
            os.makedirs(images_dir, exist_ok=True)
        except (OSError, IOError) as e:
            logger.error(f"创建输出目录失败: {e}")
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            return None
        
        # 准备文档结构
        try:
            document_structure = {
                "metadata": extract_metadata(doc, docx_path),
                "sections": [],
                "processing_info": {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "source_file": os.path.basename(docx_path),
                    "file_size_bytes": file_size,
                    "errors": [],
                    "warnings": []
                }
            }
        except Exception as e:
            logger.warning(f"提取元数据失败: {e}")
            document_structure = {
                "metadata": {"error": f"Failed to extract metadata: {e}"},
                "sections": [],
                "processing_info": {
                    "source_file": os.path.basename(docx_path),
                    "file_size_bytes": file_size,
                    "errors": [f"Metadata extraction failed: {e}"],
                    "warnings": []
                }
            }
        
        # 图片引用字典
        image_references = {}
        
        # 提取页眉页脚中的图片（增强错误处理）
        try:
            extract_header_footer_images(doc, images_dir, image_references, quick_mode)
        except Exception as e:
            logger.warning(f"提取页眉页脚图片失败，继续处理: {e}")
            document_structure["processing_info"]["warnings"].append(f"Header/footer image extraction failed: {e}")
        
        # 创建根节点
        root_section = {
            "type": "section",
            "title": "根节点",
            "level": 0,
            "content": []
        }
        document_structure["sections"].append(root_section)
        
        # 使用栈管理标题层级
        stack = deque([root_section])
        
        # 状态跟踪
        in_toc = False  # 是否在目录部分
        list_counter = defaultdict(int)  # 多级列表计数器
        table_counter = 0  # 表格计数器
        block_counter = 0  # 处理的块计数器
        
        # 遍历文档块（增强错误处理）
        try:
            for block in iter_block_items(doc):
                block_counter += 1
                try:
                    # 段落处理
                    if isinstance(block, Paragraph):
                        try:
                            text = clean_text(block.text) if block.text else ""
                        except Exception as e:
                            logger.warning(f"段落 {block_counter} 文本清理失败: {e}")
                            text = str(block.text) if block.text else ""
                        
                        # 检测目录开始
                        if not in_toc and re.match(r'^(目\s*录|contents?)$', text, re.IGNORECASE):
                            in_toc = True
                            continue
                            
                        # 检测目录结束
                        if in_toc:
                            try:
                                if get_heading_level(block) > 0 or len(text) > 50:
                                    in_toc = False
                                else:
                                    continue  # 跳过目录内容
                            except Exception as e:
                                logger.warning(f"目录检测失败，继续处理: {e}")
                                in_toc = False
                        
                        # 标题处理
                        try:
                            heading_level = get_heading_level(block)
                            if heading_level > 0:
                                # 创建新章节
                                new_section = {
                                    "type": "section",
                                    "title": text,
                                    "level": heading_level,
                                    "content": []
                                }
                                
                                # 确定父章节
                                while stack and stack[-1]["level"] >= heading_level:
                                    stack.pop()
                                
                                # 添加到适当的父章节
                                parent_section = stack[-1] if stack else root_section
                                parent_section["content"].append(new_section)
                                
                                stack.append(new_section)
                                continue
                        except Exception as e:
                            logger.warning(f"标题级别检测失败: {e}")
                        
                        # 列表项处理
                        try:
                            if is_list_item(block):
                                list_level, prefix = get_list_info(block, list_counter)
                                
                                # 提取列表项中的内容（图片和SmartArt）
                                content_nodes = extract_paragraph_content(block, output_dir, image_references, quick_mode)
                                
                                # 创建列表项节点
                                list_item = {
                                    "type": "list_item",
                                    "text": text,  # 保持原始文本，不添加prefix
                                    "level": list_level,
                                    "prefix": prefix,  # 将前缀作为独立字段存储
                                    "is_bullet": "bullet" in str(block.style.name).lower() if hasattr(block, 'style') and block.style else False
                                }
                                
                                # 添加到当前章节
                                current_section = stack[-1] if stack else root_section
                                current_section["content"].append(list_item)
                                
                                # 添加内容作为独立节点
                                if content_nodes:
                                    current_section["content"].extend(content_nodes)
                                continue
                        except Exception as e:
                            logger.warning(f"列表项处理失败: {e}")
                        
                        # 普通段落处理
                        if text or (hasattr(block, 'runs') and block.runs):
                            try:
                                # 提取段落中的内容（图片和SmartArt）
                                content_nodes = extract_paragraph_content(block, output_dir, image_references, quick_mode)
                                
                                # 添加文本段落
                                if text:
                                    para_item = {
                                        "type": "paragraph",
                                        "text": text,
                                        "bold": False,
                                        "italic": False
                                    }
                                    
                                    # 检查格式
                                    try:
                                        if hasattr(block, 'runs') and block.runs:
                                            para_item["bold"] = any(run.bold for run in block.runs if run.bold)
                                            para_item["italic"] = any(run.italic for run in block.runs if run.italic)
                                    except Exception as e:
                                        logger.debug(f"格式检查失败: {e}")
                                    
                                    # 添加到当前章节
                                    current_section = stack[-1] if stack else root_section
                                    current_section["content"].append(para_item)
                                
                                # 添加内容作为独立节点
                                if content_nodes:
                                    current_section = stack[-1] if stack else root_section
                                    current_section["content"].extend(content_nodes)
                            except Exception as e:
                                logger.warning(f"段落内容提取失败: {e}")
                                document_structure["processing_info"]["warnings"].append(f"Paragraph {block_counter} content extraction failed: {e}")
                    
                    # 表格处理
                    elif isinstance(block, Table):
                        try:
                            table_data = parse_table(block, table_counter, image_references, images_dir)
                            table_item = {
                                "type": "table",
                                "index": table_counter,
                                "rows": table_data
                            }
                            table_counter += 1
                            
                            # 添加到当前章节
                            current_section = stack[-1] if stack else root_section
                            current_section["content"].append(table_item)
                        except Exception as e:
                            logger.warning(f"表格 {table_counter} 处理失败: {e}")
                            document_structure["processing_info"]["warnings"].append(f"Table {table_counter} processing failed: {e}")
                            table_counter += 1
                
                except Exception as e:
                    logger.warning(f"处理文档块 {block_counter} 时出错: {e}")
                    document_structure["processing_info"]["warnings"].append(f"Block {block_counter} processing failed: {e}")
                    continue
                    
        except Exception as e:
            logger.error(f"遍历文档块时出现严重错误: {e}")
            document_structure["processing_info"]["errors"].append(f"Document traversal failed: {e}")
        
        # 添加页眉页脚图片处理（暂时禁用，因为API限制）
        # 注意：页眉页脚图片功能可以在需要时启用
        header_footer_images = []
        # try:
        #     extract_header_footer_images(doc, images_dir, image_references)
        # except Exception as e:
        #     logger.warning(f"页眉页脚图片处理失败: {e}")
        
        if header_footer_images:
            document_structure["header_footer_images"] = [
                {
                    "type": "image",
                    "url": img["url"],
                    "format": img["format"],
                    "width": img["width"],
                    "height": img["height"],
                    "size": img["size"],
                    "location": "header" if "header" in img.get("context", "") else "footer"
                } 
                for img in header_footer_images
            ]
        
        # 添加图片引用
        if image_references:
            document_structure["images"] = image_references
            logger.info(f"检测到 {len(image_references)} 张图片")
        else:
            logger.info(f"文档 {os.path.basename(docx_path)} 中没有检测到图片")
        
        # 添加处理统计信息
        document_structure["processing_info"]["blocks_processed"] = block_counter
        document_structure["processing_info"]["tables_found"] = table_counter
        document_structure["processing_info"]["images_found"] = len(image_references)
        
        # 清理临时目录
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
                logger.info(f"清理临时目录: {temp_dir}")
        except Exception as e:
            logger.warning(f"清理临时目录失败: {e}")
        
        return document_structure
        
    except Exception as e:
        logger.error(f"解析文档 {docx_path} 失败: {e}")
        logger.error(traceback.format_exc())
        
        # 确保清理临时目录
        try:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
        except:
            pass
            
        return None

def add_error_to_failed_files(failed_files, filename, error_msg):
    """添加错误信息到失败文件列表的辅助函数"""
    failed_files.append({"file": filename, "error": error_msg})

def safe_filename(filename):
    """创建安全的文件名，移除无效字符"""
    # 移除文件系统不允许的字符
    safe_name = re.sub(r'[<>:"/\\|?*\[\]]', '_', filename)
    # 限制长度
    return safe_name[:150]

def process_docx_folder(input_folder, output_base_dir, quick_mode=True):
    """
    批量处理文件夹中的所有DOCX文件，增强错误处理和进度跟踪
    
    Args:
        input_folder: 输入文件夹路径
        output_base_dir: 输出基础目录
        quick_mode: 快速模式，跳过耗时的EMF/WMF转换（默认True）
    """
    try:
        # 确保输出目录存在
        os.makedirs(output_base_dir, exist_ok=True)
    except (OSError, IOError) as e:
        logger.error(f"创建输出目录失败: {e}")
        return 0
    
    # 检查输入目录
    if not os.path.exists(input_folder):
        logger.error(f"输入目录不存在: {input_folder}")
        return 0
        
    if not os.path.isdir(input_folder):
        logger.error(f"输入路径不是目录: {input_folder}")
        return 0
    
    # 准备汇总数据
    all_documents = []
    
    # 查找所有DOCX文件，过滤掉临时文件
    try:
        all_files = os.listdir(input_folder)
        docx_files = [
            f for f in all_files 
            if f.lower().endswith('.docx') and not f.startswith('~$')
        ]
    except (OSError, IOError) as e:
        logger.error(f"读取目录失败: {e}")
        return 0
    
    if not docx_files:
        logger.warning(f"在 {input_folder} 中没有找到有效的DOCX文件")
        return 0
    
    logger.info(f"开始处理 {len(docx_files)} 个DOCX文件...")
    
    processed_count = 0
    failed_files = []
    skipped_files = []
    
    for idx, filename in enumerate(docx_files, 1):
        docx_path = os.path.join(input_folder, filename)
        logger.info(f"处理文件 ({idx}/{len(docx_files)}): {filename}")
        
        try:
            # 为每个文件创建输出目录
            safe_name = safe_filename(filename.replace('.docx', ''))
            output_dir = os.path.join(output_base_dir, safe_name)
            
            try:
                os.makedirs(output_dir, exist_ok=True)
            except (OSError, IOError) as e:
                logger.error(f"创建文件输出目录失败: {e}")
                add_error_to_failed_files(failed_files, filename, f"Directory creation failed: {e}")
                continue
            
            # 检查文件是否可读
            try:
                with open(docx_path, 'rb') as test_file:
                    test_file.read(1)
            except (OSError, IOError) as e:
                logger.error(f"文件不可读: {filename}, 错误: {e}")
                add_error_to_failed_files(failed_files, filename, f"File not readable: {e}")
                continue
            
            # 解析文档
            document_structure = parse_docx(docx_path, output_dir, quick_mode)
            
            if not document_structure:
                logger.error(f"跳过 {filename}，解析失败")
                add_error_to_failed_files(failed_files, filename, "Document parsing failed")
                continue
            
            # 检查解析结果的质量
            if document_structure.get("processing_info", {}).get("errors"):
                logger.warning(f"文件 {filename} 解析时有错误，但继续处理")
            
            processed_count += 1
            
            # 保存为JSON文件
            json_path = os.path.join(output_dir, "document.json")
            try:
                with open(json_path, "w", encoding="utf-8") as f:
                    json.dump(document_structure, f, ensure_ascii=False, indent=2)
            except (OSError, IOError) as e:
                logger.error(f"保存JSON文件失败: {e}")
                # 即使JSON保存失败，也算作处理失败
                processed_count -= 1
                add_error_to_failed_files(failed_files, filename, f"JSON save failed: {e}")
                continue
            
            # 添加到汇总列表
            all_documents.append({
                "file": filename,
                "path": json_path,
                "status": "success",
                "images_found": len(document_structure.get("images", {})),
                "warnings": len(document_structure.get("processing_info", {}).get("warnings", [])),
                "errors": len(document_structure.get("processing_info", {}).get("errors", []))
            })
            
            # 统计图片数量
            total_images = len(document_structure.get("images", {}))
            hf_images = len(document_structure.get("header_footer_images", []))
            processing_info = document_structure.get("processing_info", {})
            
            logger.info(f"文件 {filename} 处理完成 - 图片: {total_images}, 页眉页脚图片: {hf_images}, 警告: {len(processing_info.get('warnings', []))}")
            
        except Exception as e:
            logger.error(f"处理文件 {filename} 时发生未知错误: {e}")
            logger.error(traceback.format_exc())
            add_error_to_failed_files(failed_files, filename, f"Unexpected error: {e}")
            continue
    
    # 保存汇总信息
    try:
        summary_path = os.path.join(output_base_dir, "summary.json")
        summary_data = {
            "input_folder": input_folder,
            "processing_timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_files": len(docx_files),
            "processed": processed_count,
            "failed": len(failed_files),
            "skipped": len(skipped_files),
            "failed_files": failed_files,
            "skipped_files": skipped_files,
            "documents": all_documents,
            "success_rate": f"{processed_count/len(docx_files)*100:.1f}%" if docx_files else "0%"
        }
        
        with open(summary_path, "w", encoding="utf-8") as f:
            json.dump(summary_data, f, ensure_ascii=False, indent=2)
            
        logger.info(f"汇总信息保存在 {summary_path}")
        
    except Exception as e:
        logger.error(f"保存汇总信息失败: {e}")
    
    # 打印处理结果
    logger.info(f"批量处理完成!")
    logger.info(f"成功处理: {processed_count}/{len(docx_files)} 个文件 ({processed_count/len(docx_files)*100:.1f}%)")
    
    if failed_files:
        logger.warning(f"失败文件 ({len(failed_files)}):")
        for failed in failed_files[:5]:  # 只显示前5个失败文件
            logger.warning(f"  - {failed['file']}: {failed['error']}")
        if len(failed_files) > 5:
            logger.warning(f"  ... 还有 {len(failed_files)-5} 个失败文件，详见summary.json")
    
    return processed_count

if __name__ == "__main__":
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("用法: python docx_parser.py <输入路径> [输出路径]")
        print("示例: python docx_parser.py demo.docx")
        print("示例: python docx_parser.py demo.docx my_output/")
        print("示例: python docx_parser.py Files/PLM2.0/")
        print("示例: python docx_parser.py Files/PLM2.0/ my_batch_output/")
        sys.exit(1)
    
    input_path = sys.argv[1]
    custom_output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 检查输入路径是否存在
    if not os.path.exists(input_path):
        logger.error(f"输入路径不存在: {input_path}")
        sys.exit(1)
    
    # 判断是单个文件还是文件夹
    if os.path.isfile(input_path):
        # 处理单个文件
        if not input_path.lower().endswith('.docx'):
            logger.error(f"不是有效的DOCX文件: {input_path}")
            sys.exit(1)
        
        # 为单个文件创建输出目录
        file_basename = os.path.splitext(os.path.basename(input_path))[0]
        safe_name = safe_filename(file_basename)
        
        if custom_output_dir:
            # 使用自定义输出路径
            output_dir = os.path.join(custom_output_dir, safe_name)
        else:
            # 使用默认输出路径
            output_dir = f"parsed_docs/single_files/{safe_name}"
        
        os.makedirs(output_dir, exist_ok=True)
        
        logger.info(f"开始处理单个文件: {input_path}")
        logger.info(f"输出目录: {output_dir}")
        
        # 解析文档（使用快速模式）
        document_structure = parse_docx(input_path, output_dir, quick_mode=True)
        
        if document_structure:
            # 保存为JSON文件
            json_path = os.path.join(output_dir, "document.json")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(document_structure, f, ensure_ascii=False, indent=2)
            
            # 统计信息
            total_images = len(document_structure.get("images", {}))
            processing_info = document_structure.get("processing_info", {})
            warnings = len(processing_info.get("warnings", []))
            errors = len(processing_info.get("errors", []))
            
            logger.info(f"文件处理完成!")
            logger.info(f"输出位置: {output_dir}")
            logger.info(f"图片数量: {total_images}")
            logger.info(f"警告数量: {warnings}")
            logger.info(f"错误数量: {errors}")
            
            if warnings > 0 or errors > 0:
                logger.info("详细的警告和错误信息请查看 document.json 文件")
        else:
            logger.error("文件处理失败")
            sys.exit(1)
            
    elif os.path.isdir(input_path):
        # 批量处理文件夹
        # 创建与输入文件夹对应的输出目录
        folder_name = os.path.basename(input_path.rstrip('/'))
        
        if custom_output_dir:
            # 使用自定义输出路径
            output_base_dir = os.path.join(custom_output_dir, f"{folder_name}_parsed")
        else:
            # 使用默认输出路径
            output_base_dir = f"parsed_docs/{folder_name}_parsed_without_modular"
        
        logger.info(f"开始批量处理文件夹: {input_path}")
        logger.info(f"输出目录: {output_base_dir}")
        
        # 处理所有DOCX文件
        processed_count = process_docx_folder(input_path, output_base_dir, quick_mode=True)
        
        # 打印总结报告
        if processed_count and processed_count > 0:
            summary_path = os.path.join(output_base_dir, "summary.json")
            if os.path.exists(summary_path):
                with open(summary_path, "r", encoding="utf-8") as f:
                    summary = json.load(f)
                    print("\n" + "="*60)
                    print("📊 处理总结报告")
                    print("="*60)
                    print(f"📁 处理文件夹: {summary['input_folder']}")
                    print(f"📄 文件总数: {summary['total_files']}")
                    print(f"✅ 成功处理: {summary['processed']}")
                    print(f"❌ 失败文件: {summary['failed']}")
                    print(f"📈 成功率: {summary.get('success_rate', 'N/A')}")
                    
                    if summary.get('failed_files'):
                        print("\n❌ 失败文件列表:")
                        for failed in summary['failed_files'][:20]:
                            print(f"   - {failed.get('file', 'Unknown')}: {failed.get('error', 'Unknown error')}")
                        if len(summary['failed_files']) > 20:
                            print(f"   ... 还有 {len(summary['failed_files'])-20} 个失败文件")

                    print(f"\n📁 结果保存在: {output_base_dir}")
                    print("="*60)
        else:
            logger.error("没有成功处理任何文件")
            sys.exit(1)
    else:
        logger.error(f"无法识别的输入路径类型: {input_path}")
        sys.exit(1)
