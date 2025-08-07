#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
测试 SmartArt 检测和解析脚本
"""

import os
import zipfile
from docx import Document
from lxml import etree
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def find_smartart_in_docx(docx_path):
    """
    在DOCX文件中查找SmartArt内容
    """
    logger.info(f"分析文档: {docx_path}")
    
    # 方法1: 使用zipfile直接检查XML内容
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 列出所有文件
            file_list = docx_zip.namelist()
            logger.info(f"文档包含的文件: {len(file_list)}")
            
            # 检查是否有SmartArt相关的文件
            smartart_files = [f for f in file_list if 'diagram' in f.lower() or 'smartart' in f.lower()]
            if smartart_files:
                logger.info(f"发现SmartArt相关文件: {smartart_files}")
            
            # 检查主文档内容
            try:
                with docx_zip.open('word/document.xml') as doc_xml:
                    content = doc_xml.read().decode('utf-8')
                    
                    # 查找SmartArt相关的XML标签
                    smartart_indicators = [
                        'dgm:relIds',           # Diagram关系ID
                        'wp:docPr',             # Drawing properties
                        'a:graphic',            # Graphic element
                        'dgm:',                 # Diagram namespace
                        'wpg:',                 # Word processing group
                        'wps:',                 # Word processing shape
                        'mc:AlternateContent',  # Alternative content (often used for SmartArt)
                        'mc:Choice',            # Choice element
                        'mc:Fallback'           # Fallback element
                    ]
                    
                    found_indicators = []
                    for indicator in smartart_indicators:
                        if indicator in content:
                            found_indicators.append(indicator)
                            # 查找包含该指示器的行
                            lines = content.split('\n')
                            for i, line in enumerate(lines):
                                if indicator in line:
                                    logger.info(f"第{i+1}行发现 {indicator}: {line.strip()[:100]}...")
                                    break
                    
                    if found_indicators:
                        logger.info(f"发现SmartArt指示器: {found_indicators}")
                        return True
                    else:
                        logger.info("未发现SmartArt指示器")
                        
            except Exception as e:
                logger.error(f"读取document.xml失败: {e}")
                
    except Exception as e:
        logger.error(f"分析文档失败: {e}")
        return False
    
    # 方法2: 使用python-docx检查
    try:
        doc = Document(docx_path)
        
        # 检查段落中的drawing元素
        for para in doc.paragraphs:
            if para._element.xml:
                if any(tag in para._element.xml for tag in ['<w:drawing', '<mc:AlternateContent', '<dgm:']):
                    logger.info(f"段落中发现SmartArt: {para.text[:50]}...")
                    logger.info(f"XML内容: {para._element.xml[:200]}...")
                    return True
        
        # 检查表格中的drawing元素
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para._element.xml:
                            if any(tag in para._element.xml for tag in ['<w:drawing', '<mc:AlternateContent', '<dgm:']):
                                logger.info(f"表格单元格中发现SmartArt: {para.text[:50]}...")
                                return True
                                
    except Exception as e:
        logger.error(f"使用python-docx分析失败: {e}")
    
    return False

def extract_smartart_details(docx_path):
    """
    提取SmartArt的详细信息
    """
    logger.info(f"提取SmartArt详细信息: {docx_path}")
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 读取主文档
            with docx_zip.open('word/document.xml') as doc_xml:
                content = doc_xml.read().decode('utf-8')
                
                # 解析XML
                root = etree.fromstring(content.encode('utf-8'))
                
                # 定义命名空间
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
                    'dgm': 'http://schemas.openxmlformats.org/drawingml/2006/diagram',
                    'wpg': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingGroup',
                    'wps': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingShape'
                }
                
                # 查找drawing元素
                drawings = root.findall('.//w:drawing', namespaces)
                logger.info(f"发现 {len(drawings)} 个drawing元素")
                
                for i, drawing in enumerate(drawings):
                    logger.info(f"Drawing {i+1}:")
                    
                    # 查找graphic元素
                    graphics = drawing.findall('.//a:graphic', namespaces)
                    for j, graphic in enumerate(graphics):
                        graphic_data = graphic.find('.//a:graphicData', namespaces)
                        if graphic_data is not None:
                            uri = graphic_data.get('uri', '')
                            logger.info(f"  Graphic {j+1} URI: {uri}")
                            
                            # 检查是否为SmartArt (diagram)
                            if 'diagram' in uri:
                                logger.info(f"  >>> 这是一个SmartArt图表!")
                                # 查找关系ID
                                rel_ids = graphic_data.findall('.//dgm:relIds', namespaces)
                                for rel_id in rel_ids:
                                    logger.info(f"    关系ID: {etree.tostring(rel_id, encoding='unicode')}")
                
                # 查找AlternateContent元素（通常包含SmartArt的兼容性内容）
                alt_contents = root.findall('.//mc:AlternateContent', namespaces)
                logger.info(f"发现 {len(alt_contents)} 个AlternateContent元素")
                
                for i, alt_content in enumerate(alt_contents):
                    logger.info(f"AlternateContent {i+1}:")
                    
                    choices = alt_content.findall('.//mc:Choice', namespaces)
                    for choice in choices:
                        requires = choice.get('Requires', '')
                        logger.info(f"  Choice requires: {requires}")
                        
                        # 检查choice中的内容
                        choice_drawings = choice.findall('.//w:drawing', namespaces)
                        if choice_drawings:
                            logger.info(f"    包含 {len(choice_drawings)} 个drawing元素")
                    
                    fallbacks = alt_content.findall('.//mc:Fallback', namespaces)
                    for fallback in fallbacks:
                        logger.info(f"  Fallback内容存在")
                        fallback_drawings = fallback.findall('.//w:drawing', namespaces)
                        if fallback_drawings:
                            logger.info(f"    包含 {len(fallback_drawings)} 个drawing元素")
                
    except Exception as e:
        logger.error(f"提取SmartArt详细信息失败: {e}")

def main():
    """主函数"""
    # 测试文件夹中的一些文档
    test_files = [
        "Files/PLM2.0/BOM 审核申请.docx",
        "Files/PLM2.0/包装标签设计申请流程.docx", 
        "Files/PLM2.0/MCT流程.docx",
        "Files/PLM2.0/制造商申请V2.0.docx"
    ]
    
    for docx_file in test_files:
        if os.path.exists(docx_file):
            logger.info(f"\n{'='*60}")
            logger.info(f"分析文件: {docx_file}")
            logger.info(f"{'='*60}")
            
            # 检查是否包含SmartArt
            has_smartart = find_smartart_in_docx(docx_file)
            
            if has_smartart:
                logger.info("该文档包含SmartArt内容，提取详细信息...")
                extract_smartart_details(docx_file)
            else:
                logger.info("该文档不包含SmartArt内容")
        else:
            logger.warning(f"文件不存在: {docx_file}")

if __name__ == "__main__":
    main()
